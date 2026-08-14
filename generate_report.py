import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Colors
    # Primary: Dark Navy (#1E293B), Secondary: Indigo (#4F46E5), Accent: Emerald (#10B981)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Zynovix Tech Solution")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 41, 59)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Website Changes & Light/Dark Theme Implementation Report")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(15)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(79, 70, 229)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("Date: August 13, 2026  |  Branch: main  |  Project: zynovixtechsolution")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spacer

    # Section 1: Executive Summary
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Executive Summary (मॉस रिपोर्ट सारांश)")
    r_h1.font.name = "Arial"
    r_h1.font.color.rgb = RGBColor(30, 41, 59)

    p_exec = doc.add_paragraph()
    p_exec.add_run(
        "इस प्रोजेक्ट में संपूर्ण वेबसाइट पर Light Theme & Dark Theme System का सफल क्रियान्वयन (Implementation) किया गया है। "
        "वेबसाइट के सभी 26 कंपोनेंट्स और पेजों को एडिट करके हाई-कंट्रास्ट Light Mode और Dark Mode संगत बनाया गया है। "
        "इसके अतिरिक्त, यूजर एक्सपीरियंस (UX) को बेहतर बनाने के लिए फ्लोटिंग Theme Switcher बटन और WhatsApp Quick Support Widget भी जोड़े गए हैं।"
    )

    # Highlights Box
    p_stat = doc.add_paragraph()
    p_stat.add_run("• कुल मॉडिफाइड फाइल्स (Modified Files): 26\n").bold = True
    p_stat.add_run("• नई जोड़ी गई फाइल्स (New Untracked Files): 4\n").bold = True
    p_stat.add_run("• कुल कोड बदलाव (Code Changes): +364 Insertions, -401 Deletions\n").bold = True
    p_stat.add_run("• मुख्य फ़ीचर्स (Key Features): ThemeProvider, ThemeToggle (Sun/Moon), WhatsApp Button, CSS Theme Variables").bold = True

    doc.add_paragraph()

    # Section 2: New Files Created
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. New Files Added (नये जोड़े गए कंपोनेंट्स)")
    r_h2.font.name = "Arial"
    r_h2.font.color.rgb = RGBColor(30, 41, 59)

    table1 = doc.add_table(rows=1, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table1.rows[0].cells
    hdr_titles = ["File Name", "Location", "Description & Purpose"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)

    new_files_data = [
        ("ThemeToggle.tsx", "components/ThemeToggle.tsx", "Header/Navbar में Sun ☀️ / Moon 🌙 आइकन वाला Theme Switcher बटन"),
        ("ThemeProvider.tsx", "components/ThemeProvider.tsx", "Next.js Theme Context Provider (next-themes wrapper)"),
        ("ThemeControls.tsx", "components/ThemeControls.tsx", "Theme customization एवं control settings component"),
        ("WhatsAppButton.tsx", "components/WhatsAppButton.tsx", "वेबसाइट पर डायरेक्ट कस्टमर कनेक्ट के लिए Floating WhatsApp Widget")
    ]

    for filename, loc, desc in new_files_data:
        row_cells = table1.add_row().cells
        row_cells[0].text = filename
        row_cells[1].text = loc
        row_cells[2].text = desc
        for cell in row_cells:
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, 120, 120, 150, 150)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = "Arial"

    doc.add_paragraph()

    # Section 3: Modified Pages (app/)
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Modified Pages (app/ Directory)")
    r_h3.font.name = "Arial"
    r_h3.font.color.rgb = RGBColor(30, 41, 59)

    table2 = doc.add_table(rows=1, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells2 = table2.rows[0].cells
    hdr_titles2 = ["Page Path", "Changes Applied", "Impact"]
    for i, title in enumerate(hdr_titles2):
        hdr_cells2[i].text = title
        set_cell_background(hdr_cells2[i], "4F46E5")
        for p in hdr_cells2[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)

    pages_data = [
        ("app/layout.tsx", "ThemeProvider & WhatsAppButton को Global Root Layout में wrap किया गया", "High"),
        ("app/globals.css", "Light/Dark mode के लिए CSS Color Variables और Theme Classes जोड़े गए", "Critical"),
        ("app/about/page.tsx", "About page के text contrast, mission boxes, hero, team sections में light mode styling", "Medium"),
        ("app/contact/page.tsx", "Contact form, input boxes, labels, submit button & contact cards styling", "High"),
        ("app/services/page.tsx", "Services listing cards, icons & hover backgrounds for Light/Dark mode", "Medium"),
        ("app/portfolio/page.tsx", "Portfolio project grid, filter buttons & tag badges compatibility", "Medium"),
        ("app/team/page.tsx", "Team member cards, social links & overlays theme adaptation", "Medium"),
        ("app/careers/page.tsx", "Career page container & section responsiveness for dual theme", "Medium")
    ]

    for p_path, changes, impact in pages_data:
        row_cells = table2.add_row().cells
        row_cells[0].text = p_path
        row_cells[1].text = changes
        row_cells[2].text = impact
        for cell in row_cells:
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, 100, 100, 120, 120)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = "Arial"

    doc.add_paragraph()

    # Section 4: Modified Components (components/)
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Modified UI Components (components/ Directory)")
    r_h4.font.name = "Arial"
    r_h4.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_heading("A. Layout & Navigation", level=2)
    p_nav = doc.add_paragraph()
    p_nav.add_run("• components/layout/Navbar.tsx: ").bold = True
    p_nav.add_run("ThemeToggle Switcher जोड़ा गया, mobile menu, text colors और background blur update किए गए।\n")
    p_nav.add_run("• components/layout/Footer.tsx: ").bold = True
    p_nav.add_run("Footer contrast, links, social icons और muted text color Light mode के अनुसार एडजस्ट किए गए।")

    doc.add_heading("B. Home Sections", level=2)
    sections_list = [
        ("Hero.tsx", "Hero section titles, subheadings, CTA buttons, backdrops और ambient lighting fix"),
        ("AboutPreview.tsx", "Preview card styling, typography contrast for light mode"),
        ("ServicesPreview.tsx", "Dynamic light/dark cards for service previews"),
        ("PortfolioPreview.tsx", "Project cards preview styling & hover states"),
        ("CTA.tsx", "Call-to-action banner background, text contrast & button colors"),
        ("Stats.tsx", "Statistic counters visibility & borders fix"),
        ("TrustBar.tsx", "Partner/Tech stack logos border & background adjustment")
    ]
    for filename, desc in sections_list:
        p_sec = doc.add_paragraph()
        p_sec.add_run(f"• components/sections/{filename}: ").bold = True
        p_sec.add_run(desc)

    doc.add_heading("C. Careers Components", level=2)
    careers_list = [
        ("ApplicationForm.tsx", "Job application form inputs, labels, upload buttons & validation styling"),
        ("OpenPositions.tsx", "Job listing search input, position cards & category badges"),
        ("Benefits.tsx", "Perks/Benefits cards light mode background & icon contrast"),
        ("FAQ.tsx", "FAQ accordion background, border & active tab text colors"),
        ("HiringProcess.tsx", "Process timeline step cards & connector lines styling"),
        ("WhyJoin.tsx & CareerHero.tsx & CTA.tsx", "Career sections visual theme alignment")
    ]
    for filename, desc in careers_list:
        p_car = doc.add_paragraph()
        p_car.add_run(f"• components/careers/{filename}: ").bold = True
        p_car.add_run(desc)

    doc.add_paragraph()

    # Section 5: Conclusion & Recommendation
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. Conclusion & Verification (निष्कर्ष एवं सत्यापन)")
    r_h5.font.name = "Arial"
    r_h5.font.color.rgb = RGBColor(30, 41, 59)

    p_conc = doc.add_paragraph()
    p_conc.add_run(
        "1. सभी मॉडिफाइड फाइल्स सफलतापूर्वक अपडेट कर दी गई हैं और `npm run dev` के जरिए टेस्ट की जा चुकी हैं।\n"
        "2. Light Theme और Dark Theme दोनों मोड में वेबसाइट बिना किसी लेआउट डिस्टॉर्शन के उत्कृष्ट (Premium UI) दिखती है।\n"
        "3. वेबसाइट पर व्हाट्सएप चैट सपोर्ट बटन लाइव काम कर रहा है।"
    )

    # Save output
    output_docx = r"f:\Office_Project\Praveen\zynovixtechsolution - Copy\Website_Changes_Report.docx"
    doc.save(output_docx)
    print(f"Report successfully saved to: {output_docx}")

if __name__ == "__main__":
    create_report()
